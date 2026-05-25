# -*- coding: utf-8 -*-
"""
Genera DOC-PEDIDOS-MAYORISTA-001 — Documentación completa en Word.
Uso: python docs/scripts/generar_documentacion_word.py
Requiere: pip install python-docx
"""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Instale python-docx: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "DOC-PEDIDOS-MAYORISTA-001-Documentacion-Completa.docx"
HEADER_FILL = "8B1A1A"


def shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def style_table_header(table, cols: int):
    for i in range(cols):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)


def add_table(doc, headers, rows, col_widths_cm=None):
    ncol = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows):
        cells = list(row) + [""] * (ncol - len(row))
        for j in range(ncol):
            t.rows[i + 1].cells[j].text = str(cells[j])[:800]
    style_table_header(t, ncol)
    if col_widths_cm:
        for row in t.rows:
            for j, w in enumerate(col_widths_cm):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return t


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def p(doc, text):
    doc.add_paragraph(text)


def bullet(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SUPERTIENDAS CAÑAVERAL\n")
    r.font.size = Pt(14)
    r.font.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("DOCUMENTACIÓN COMPLETA DEL SISTEMA\n")
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(139, 26, 26)
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "Gestión de Pedidos Mayorista — Carnicería\n"
        "Pedidos en tiempo real, catálogo de cortes y operación multi-sede"
    )
    r3.font.size = Pt(12)
    doc.add_paragraph()
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Código del documento", "DOC-PEDIDOS-MAYORISTA-001"],
            ["Nombre del producto", "Pedidos Mayorista (Supertiendas Cañaveral)"],
            ["Versión documental", "1.0"],
            ["Fecha de elaboración", str(date.today())],
            ["Clasificación", "Uso interno — Confidencial"],
            ["API", "Supertiendas Cañaveral API (FastAPI)"],
            ["Repositorio / carpeta", r"D:\Pedidos mayorista"],
        ],
        [5, 11],
    )
    doc.add_page_break()

    h1(doc, "Control del documento")
    h2(doc, "Historial de revisiones")
    add_table(
        doc,
        ["Versión", "Fecha", "Descripción"],
        [
            ["1.0", "2026-05-25", "Documentación inicial del sistema de pedidos mayorista"],
            ["1.0", str(date.today()), "Exportación a formato Word con tablas e ISO"],
        ],
    )
    h2(doc, "Distribución y aprobación")
    add_table(
        doc,
        ["Rol", "Nombre (completar)", "Firma", "Fecha"],
        [
            ["Elaboró — Desarrollo", "", "", ""],
            ["Revisó — Operación carnicería", "", "", ""],
            ["Revisó — TI", "", "", ""],
            ["Aprobó — Responsable tratamiento de datos", "", "", ""],
        ],
    )
    doc.add_page_break()

    h1(doc, "Tabla de contenidos")
    for line in [
        "1. Identificación y alcance",
        "2. Contexto de negocio",
        "3. Roles y permisos",
        "4. Requisitos funcionales",
        "5. Requisitos no funcionales",
        "6. Arquitectura del sistema",
        "7. Módulos frontend (pantallas)",
        "8. API REST y WebSockets",
        "9. Modelo de datos",
        "10. Catálogo de cortes (res)",
        "11. Flujos de negocio",
        "12. Seguridad",
        "13. Operación, despliegue y scripts",
        "14. Matriz ISO y normatividad (Colombia)",
        "15. Brechas y plan de cierre",
        "16. Glosario y anexos",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # 1
    h1(doc, "1. Identificación y alcance")
    h2(doc, "1.1 Objeto")
    p(
        doc,
        "Sistema web para la gestión de pedidos de carnicería en tiempo real entre el área mayorista "
        "y las sedes (carnicería). Permite crear pedidos con cortes y cantidades en kilogramos, "
        "notificar a la sede vía WebSocket, asignar carnicero, cambiar estados (pendiente, en proceso, "
        "finalizado), supervisión por jefe de carnes y administración central (usuarios, sedes, catálogo).",
    )
    h2(doc, "1.2 Alcance incluido")
    bullet(
        doc,
        [
            "Login por rol con redirección automática (mayorista, sede, jefe carnes, admin).",
            "Creación de pedidos con detalle por corte, tipo de corte y kg.",
            "Numeración consecutiva de pedido por sede (sin reinicio diario).",
            "Tiempo real: new_order, order_update, order_problem, availability_update.",
            "Catálogo de cortes de res en PostgreSQL + imágenes en servidor (/static/cortes/res/).",
            "Gestión de carniceros, disponibilidad diaria (jefe de carnes).",
            "Panel admin: sedes, categorías, cortes, tipos de corte, usuarios, estadísticas.",
            "Interfaz responsiva (móvil, tablet, escritorio).",
            "Contraseñas con hash PBKDF2-SHA256 (Passlib); sesión JWT.",
        ],
    )
    h2(doc, "1.3 Alcance excluido")
    add_table(
        doc,
        ["Excluido", "Motivo"],
        [
            ["Facturación / POS / inventario contable", "Fuera del alcance del MVP"],
            ["Integración ERP o SAP", "No implementada"],
            ["App móvil nativa (iOS/Android)", "Solo aplicación web PWA/navegador"],
            ["Certificación ISO de la organización", "Requiere SGI corporativo aparte"],
            ["ProGest (activos fijos Cañaveral)", "Proyecto distinto en otro repositorio"],
        ],
    )

    # 2
    h1(doc, "2. Contexto de negocio")
    p(
        doc,
        "En operación de carnicería, el mayorista recibe pedidos de clientes y los transmite a la sede "
        "para preparación. Antes, la comunicación podía ser verbal o por canales lentos. Este sistema "
        "centraliza el pedido, lo numera, lo muestra en pantalla de carnicería al instante y permite "
        "seguimiento de tiempos (started_at, finished_at) y reporte de problemas.",
    )
    add_table(
        doc,
        ["Actor", "Rol técnico", "Necesidad"],
        [
            ["Mayorista", "mayorista", "Armar pedido, elegir cortes, ver historial de su sede"],
            ["Tablet / pantalla sede", "sede_butcher", "Ver pedidos entrantes, asignar carnicero, cambiar estado"],
            ["Carnicero", "carnicero", "Operar en sede (misma UI que sede)"],
            ["Jefe de carnes", "jefe_carnes", "Monitor, historial, disponibilidad de personal"],
            ["Administrador TI", "admin", "ABM sedes, productos, usuarios, estadísticas"],
        ],
    )

    # 3 Roles
    h1(doc, "3. Roles y permisos")
    add_table(
        doc,
        ["Rol (user.role)", "Ruta React", "Acceso"],
        [
            ["mayorista", "/mayorista", "Crear pedidos, catálogo, historial sede propia"],
            ["sede_butcher", "/sede", "Cola de pedidos, estados, asignación carnicero"],
            ["carnicero", "/sede", "Igual que sede_butcher"],
            ["jefe_carnes", "/jefe", "Vista multi-pedido, disponibilidad carniceros"],
            ["admin", "/admin", "Configuración global y reportes"],
        ],
    )
    p(doc, "Tras login exitoso, HomeRedirect envía al usuario a la ruta según su rol. Rutas protegidas con ProtectedRoute y AuthContext.")

    # 4 RF
    h1(doc, "4. Requisitos funcionales")
    p(doc, "Leyenda: I = Implementado, P = Parcial, N = No implementado.")
    h2(doc, "4.1 Autenticación")
    add_table(
        doc,
        ["ID", "Requisito", "Estado"],
        [
            ["RF-AUTH-01", "Login usuario/contraseña → JWT", "I"],
            ["RF-AUTH-02", "Registro de usuarios (POST /register)", "I"],
            ["RF-AUTH-03", "Logout sede (revoca session_active)", "I"],
            ["RF-AUTH-04", "Redirección por rol", "I"],
            ["RF-AUTH-05", "Recuperación contraseña por email", "N"],
        ],
    )
    h2(doc, "4.2 Pedidos")
    add_table(
        doc,
        ["ID", "Requisito", "Estado"],
        [
            ["RF-PED-01", "Crear pedido con cliente, ítems (corte, tipo, kg)", "I"],
            ["RF-PED-02", "numero_pedido consecutivo por sede (sin reinicio diario)", "I"],
            ["RF-PED-03", "Estados: pendiente, en_proceso, finalizado", "I"],
            ["RF-PED-04", "Asignar carnicero al pasar a en_proceso", "I"],
            ["RF-PED-05", "Reportar problema en pedido", "I"],
            ["RF-PED-06", "Notificación tiempo real al crear/actualizar", "I"],
            ["RF-PED-07", "Filtrar pedidos por sede_id", "I"],
            ["RF-PED-08", "Timestamps started_at / finished_at", "I"],
        ],
    )
    h2(doc, "4.3 Catálogo y admin")
    add_table(
        doc,
        ["ID", "Requisito", "Estado"],
        [
            ["RF-CAT-01", "CRUD categorías y cortes", "I"],
            ["RF-CAT-02", "CRUD tipos de corte (fileteado, entero, etc.)", "I"],
            ["RF-CAT-03", "Imágenes cortes res en servidor", "I"],
            ["RF-CAT-04", "Sincronización catálogo res al arrancar API", "I"],
            ["RF-CAT-05", "Popularidad categorías (background task)", "I"],
            ["RF-ADM-01", "CRUD sedes (crea usuario tablet sede)", "I"],
            ["RF-ADM-02", "CRUD carniceros por sede", "I"],
            ["RF-ADM-03", "Disponibilidad carniceros por fecha (jefe)", "I"],
            ["RF-ADM-04", "Estadísticas pedidos por sede y top cortes", "I"],
        ],
    )

    # 5 RNF
    h1(doc, "5. Requisitos no funcionales (ISO/IEC 25010)")
    add_table(
        doc,
        ["ID", "Característica", "Requisito", "Implementación"],
        [
            ["RNF-01", "Seguridad", "Hash PBKDF2 contraseñas", "auth.py — password_hash"],
            ["RNF-02", "Seguridad", "JWT HS256", "python-jose, SECRET_KEY"],
            ["RNF-03", "Rendimiento", "Notificaciones < 1s vía WebSocket", "Socket.IO"],
            ["RNF-04", "Usabilidad", "UI responsive", "responsive.css, CSS Modules"],
            ["RNF-05", "Compatibilidad", "API REST JSON + OpenAPI", "/docs"],
            ["RNF-06", "Mantenibilidad", "CRUD centralizado", "crud.py"],
            ["RNF-07", "Fiabilidad", "create_all tablas al arranque", "SQLAlchemy"],
            ["RNF-08", "Portabilidad", "Variables .env", "DB_*, VITE_*"],
        ],
    )

    # 6 Arquitectura
    h1(doc, "6. Arquitectura del sistema")
    p(
        doc,
        "[Navegador React] ──HTTP JWT──▶ [FastAPI + Socket.IO ASGI] ──SQLAlchemy──▶ [PostgreSQL]\n"
        "                              └── /static/cortes/res/ (imágenes)"
    )
    add_table(
        doc,
        ["Capa", "Tecnología"],
        [
            ["Frontend", "React 19, Vite, React Router, CSS Modules"],
            ["Backend", "FastAPI, Uvicorn, SQLAlchemy 2"],
            ["Tiempo real", "python-socketio (ASGI)"],
            ["Base de datos", "PostgreSQL 12+"],
            ["Auth", "JWT + Passlib PBKDF2-SHA256"],
            ["Imágenes", "Pillow (generación local), estáticos FastAPI"],
        ],
    )
    h2(doc, "6.1 Estructura de directorios")
    add_table(
        doc,
        ["Ruta", "Descripción"],
        [
            ["backend/app/main.py", "Rutas API, Socket.IO, montaje /static"],
            ["backend/app/models.py", "Entidades ORM"],
            ["backend/app/crud.py", "Lógica negocio y numeración pedidos"],
            ["backend/app/auth.py", "Hash y JWT"],
            ["backend/app/catalogo_res.py", "Catálogo cortes de res"],
            ["backend/static/cortes/res/", "Imágenes JPEG de productos"],
            ["frontend/src/pages/", "Login, Mayorista, Sede, JefeCarnes, Admin"],
            ["frontend/src/context/AuthContext.jsx", "Sesión y token"],
            ["docs/", "Documentación y este Word"],
        ],
    )

    # 7 Frontend
    h1(doc, "7. Módulos frontend (pantallas)")
    h2(doc, "7.1 Login (/login)")
    p(doc, "Formulario usuario/contraseña; guarda token y user en contexto; redirige según rol.")
    h2(doc, "7.2 Mayorista (/mayorista)")
    p(doc, "Selección de categorías y cortes con imágenes, carrito de líneas (kg, tipo de corte), nombre cliente, envío POST /pedidos. Historial de pedidos de la sede del usuario. Conexión WebSocket a sala sede_{id}.")
    h2(doc, "7.3 Sede (/sede)")
    p(doc, "Cola de pedidos pendientes y en proceso; asignación de carnicero; botones para iniciar y finalizar; escucha new_order y order_update.")
    h2(doc, "7.4 Jefe de carnes (/jefe)")
    p(doc, "Monitor de pedidos activos e historial; gestión de disponibilidad de carniceros por fecha; tabla con scroll horizontal en móvil.")
    h2(doc, "7.5 Admin (/admin)")
    p(doc, "Pestañas: sedes, categorías, cortes, tipos de corte, usuarios, carniceros, estadísticas (gráficos pedidos por sede y cortes más pedidos).")

    # 8 API
    h1(doc, "8. API REST y WebSockets")
    h2(doc, "8.1 Autenticación y usuarios")
    add_table(
        doc,
        ["Método", "Ruta", "Descripción"],
        [
            ["POST", "/login", "Login → access_token + user"],
            ["POST", "/register", "Registro usuario"],
            ["POST", "/logout", "Logout (user_id); sede: session_active=0"],
            ["GET", "/users", "Listar usuarios"],
            ["PUT", "/users/{id}", "Actualizar usuario"],
            ["DELETE", "/users/{id}", "Eliminar usuario"],
        ],
    )
    h2(doc, "8.2 Sedes, catálogo, pedidos")
    add_table(
        doc,
        ["Método", "Ruta", "Descripción"],
        [
            ["GET/POST/PUT/DELETE", "/sedes", "CRUD sedes (+ usuario tablet al crear)"],
            ["GET/POST/PUT/DELETE", "/categorias", "CRUD categorías"],
            ["GET/POST/PUT/DELETE", "/cortes", "CRUD cortes (?categoria_id)"],
            ["GET/POST/PUT/DELETE", "/tipos-corte", "CRUD tipos de corte"],
            ["GET", "/pedidos?sede_id=", "Listar pedidos"],
            ["POST", "/pedidos", "Crear pedido + emit new_order"],
            ["PUT", "/pedidos/{id}/estado", "Cambiar estado + order_update"],
            ["PUT", "/pedidos/{id}/problema", "Reportar problema + order_problem"],
        ],
    )
    h2(doc, "8.3 Carniceros y disponibilidad")
    add_table(
        doc,
        ["Método", "Ruta", "Descripción"],
        [
            ["GET", "/users/carniceros/{sede_id}", "Carniceros de la sede"],
            ["POST", "/users/carniceros", "Alta carnicero"],
            ["PUT", "/users/carniceros/{id}/availability", "Flag is_available"],
            ["GET", "/butchers/sede/{sede_id}", "Listar carniceros"],
            ["GET", "/availability/{sede_id}/{date}", "Disponibilidad por fecha"],
            ["POST", "/availability", "Actualización masiva + availability_update"],
        ],
    )
    h2(doc, "8.4 Estadísticas")
    add_table(
        doc,
        ["Método", "Ruta", "Respuesta"],
        [
            ["GET", "/stats/orders-by-sede", "Conteo pedidos por sede"],
            ["GET", "/stats/top-cuts", "Kg totales por corte"],
        ],
    )
    h2(doc, "8.5 Eventos WebSocket (Socket.IO)")
    add_table(
        doc,
        ["Evento", "Dirección", "Cuándo"],
        [
            ["join_room", "Cliente → servidor", "Unirse a sala sede_{sede_id}"],
            ["new_order", "Servidor → sala", "POST /pedidos exitoso"],
            ["order_update", "Servidor → sala", "Cambio de estado pedido"],
            ["order_problem", "Servidor → sala", "Reporte de problema"],
            ["availability_update", "Servidor → sala", "Jefe actualiza disponibilidad"],
            ["sede_logout", "Servidor → sala", "Logout tablet sede"],
        ],
    )
    p(doc, "Documentación interactiva: http://localhost:8000/docs")

    # 9 Modelo datos
    h1(doc, "9. Modelo de datos")
    add_table(
        doc,
        ["Tabla", "Descripción", "Campos clave"],
        [
            ["sedes", "Puntos de venta", "nombre, ciudad"],
            ["users", "Usuarios del sistema", "username, password_hash, role, sede_id, is_available"],
            ["categorias", "Grupos de producto", "nombre, imagen_url, popularidad_score"],
            ["cortes", "Cortes de carne", "nombre, categoria_id, imagen_url"],
            ["tipos_corte", "Forma de corte", "nombre (único)"],
            ["corte_tipocorte", "N:M corte ↔ tipo", "corte_id, tipo_corte_id"],
            ["pedidos", "Encabezado pedido", "numero_pedido, estado, mayorista_id, carnicero_id, sede_id"],
            ["detalle_pedidos", "Líneas del pedido", "corte_id, tipo_corte_id, cantidad_kg"],
            ["butcher_availability", "Disponibilidad diaria", "butcher_id, date, is_available"],
        ],
    )
    h2(doc, "9.1 Estados de pedido (enum)")
    add_table(
        doc,
        ["Estado", "Significado"],
        [
            ["pendiente", "Recién creado, esperando en sede"],
            ["en_proceso", "Carnicero asignado, en preparación"],
            ["finalizado", "Pedido completado"],
        ],
    )
    h2(doc, "9.2 Roles de usuario (enum)")
    add_table(
        doc,
        ["Rol", "Valor BD"],
        [
            ["Administrador", "admin"],
            ["Mayorista", "mayorista"],
            ["Carnicero", "carnicero"],
            ["Jefe de carnes", "jefe_carnes"],
            ["Tablet sede", "sede_butcher"],
        ],
    )

    # 10 Catálogo res
    h1(doc, "10. Catálogo de cortes (res)")
    p(
        doc,
        "Definido en backend/app/catalogo_res.py. Al iniciar el API se ejecuta ensure_cortes_res y "
        "migrar_cortes_res_existentes_a_local. Imágenes en backend/static/cortes/res/ servidas en "
        "/static/cortes/res/{archivo}. Script descargar_imagenes_res.py para regenerar desde Wikimedia o Pillow.",
    )
    add_table(
        doc,
        ["Corte", "Archivo imagen"],
        [
            ["Pecho de res", "pecho-de-res.jpg"],
            ["Costilla de res", "costilla-de-res.jpg"],
            ["Punta de anca", "punta-de-anca.jpg"],
            ["Falda de res", "falda-de-res.jpg"],
            ["Molida de res", "molida-de-res.jpg"],
            ["Sobrebarriga", "sobrebarriga.jpg"],
            ["Cola de res", "cola-de-res.jpg"],
            ["Chuleta de res", "chuleta-de-res.jpg"],
            ["Hígado de res", "higado-de-res.jpg"],
            ["Hueso de res", "hueso-de-res.jpg"],
            ["Posta negra", "posta-negra.jpg"],
            ["Muchacho redondo", "muchacho-redondo.jpg"],
            ["Lomo", "lomo.jpg"],
            ["Picaña", "picana.jpg"],
        ],
    )

    # 11 Flujos
    h1(doc, "11. Flujos de negocio")
    h2(doc, "11.1 Crear pedido (mayorista → sede)")
    add_table(
        doc,
        ["Paso", "Actor", "Acción", "Sistema"],
        [
            ["1", "Mayorista", "Selecciona cortes y kg en /mayorista", "UI carrito"],
            ["2", "Mayorista", "Ingresa nombre cliente y confirma", "POST /pedidos"],
            ["3", "Sistema", "Asigna numero_pedido (max sede + 1)", "crud._next_numero_pedido"],
            ["4", "Sistema", "Emite new_order a sala sede_{id}", "Socket.IO"],
            ["5", "Sede", "Ve pedido en cola pendiente", "UI /sede"],
        ],
    )
    h2(doc, "11.2 Preparación en sede")
    add_table(
        doc,
        ["Paso", "Actor", "Acción", "Sistema"],
        [
            ["1", "Sede", "Asigna carnicero", "PUT estado=en_proceso"],
            ["2", "Sistema", "Registra started_at", "crud.update_pedido_estado"],
            ["3", "Sede", "Marca finalizado", "PUT estado=finalizado, finished_at"],
            ["4", "Sistema", "Emite order_update", "Socket.IO"],
        ],
    )
    h2(doc, "11.3 Numeración de pedidos")
    p(
        doc,
        "Regla de negocio: el campo numero_pedido es un consecutivo entero por sede (1, 2, 3…). "
        "NO se reinicia al cambiar el día calendario. Cada sede mantiene su propia secuencia. "
        "Implementación: _next_numero_pedido en crud.py consulta el máximo histórico de la sede.",
    )

    # 12 Seguridad
    h1(doc, "12. Seguridad")
    add_table(
        doc,
        ["Control", "Implementación", "Recomendación producción"],
        [
            ["Contraseñas", "password_hash PBKDF2-SHA256", "Política de complejidad corporativa"],
            ["Sesión", "JWT Bearer", "SECRET_KEY fuerte y rotación"],
            ["CORS", "allow_origins * en código", "Restringir dominios en producción"],
            ["API sin rate limit", "No implementado", "Añadir límite por IP/usuario"],
            ["HTTPS", "Responsabilidad despliegue", "Obligatorio en producción"],
        ],
    )
    p(doc, "BRECHA: no hay auditoría de intentos fallidos de login ni bandeja de alertas (a diferencia de ProGest).")

    # 13 Operación
    h1(doc, "13. Operación, despliegue y scripts")
    h2(doc, "13.1 Desarrollo local")
    add_table(
        doc,
        ["Paso", "Comando / URL"],
        [
            ["BD", "CREATE DATABASE supertiendas_db; configurar backend/.env"],
            ["Datos prueba", "python setup_initial_data.py"],
            ["API", "uvicorn app.main:app --reload --port 8000"],
            ["Front", "cd frontend && npm run dev → http://localhost:5173"],
            ["Swagger", "http://localhost:8000/docs"],
        ],
    )
    h2(doc, "13.2 Variables de entorno")
    add_table(
        doc,
        ["Variable", "Ubicación", "Descripción"],
        [
            ["DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASS", "backend/.env", "PostgreSQL"],
            ["SECRET_KEY", "backend/.env", "Firma JWT"],
            ["ACCESS_TOKEN_EXPIRE_MINUTES", "backend/.env", "Duración token (ej. 480)"],
            ["VITE_API_URL", "frontend/.env", "URL del API"],
            ["VITE_WS_URL", "frontend/.env", "URL WebSocket (mismo host API)"],
            ["PUBLIC_API_URL", "backend (opc.)", "Base URL pública para imágenes"],
        ],
    )
    h2(doc, "13.3 Scripts de mantenimiento")
    add_table(
        doc,
        ["Script", "Uso", "Riesgo"],
        [
            ["setup_initial_data.py", "Sede, categoría Res, mayorista_test", "Seguro"],
            ["create_admin.py", "Usuario administrador", "Seguro"],
            ["seed_cortes_res_servidor.py", "Insertar/actualizar cortes", "Seguro"],
            ["descargar_imagenes_res.py", "Imágenes locales en static/", "Seguro"],
            ["reset_db.py", "Borra y recrea tablas", "DESTRUCTIVO"],
            ["migrate_*.py", "Migraciones puntuales", "Revisar antes de ejecutar"],
        ],
    )
    h2(doc, "13.4 Usuario de prueba")
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Usuario", "mayorista_test"],
            ["Contraseña", "test123"],
            ["Condición", "Tras ejecutar setup_initial_data.py"],
        ],
    )

    # 14 ISO
    h1(doc, "14. Matriz de cumplimiento ISO y normatividad")
    p(doc, "Leyenda: Cumple | Parcial | No cumple | N/A. No constituye certificación ISO ni dictamen legal.")
    h2(doc, "14.1 ISO/IEC 12207 — Ciclo de vida")
    add_table(
        doc,
        ["Proceso", "Estado", "Evidencia", "Brecha"],
        [
            ["Requisitos", "Parcial", "README.md + este documento", "SRS formal con IDs trazables"],
            ["Diseño", "Parcial", "models.py, App.jsx", "Diagramas C4 / ADR"],
            ["Implementación", "Cumple", "Código backend y frontend", "—"],
            ["Pruebas", "No cumple", "Manual", "Plan ISO 29119, pytest, E2E"],
            ["Despliegue", "Parcial", "README", "CI/CD documentado"],
        ],
    )
    h2(doc, "14.2 ISO/IEC 25010 — Calidad")
    add_table(
        doc,
        ["Característica", "Estado", "Observación"],
        [
            ["Funcionalidad", "Cumple", "Flujo pedidos completo"],
            ["Rendimiento", "Parcial", "Sin pruebas de carga documentadas"],
            ["Usabilidad", "Parcial", "Responsive; sin UAT formal"],
            ["Seguridad", "Parcial", "Hash + JWT; CORS abierto"],
            ["Mantenibilidad", "Cumple", "crud.py centralizado"],
        ],
    )
    h2(doc, "14.3 ISO/IEC 27001 (referencia)")
    add_table(
        doc,
        ["Control", "Estado", "Nota"],
        [
            ["Política de seguridad", "No cumple", "Documento corporativo pendiente"],
            ["Control de acceso", "Parcial", "JWT; validación por rol en frontend principalmente"],
            ["Registro de eventos", "No cumple", "Sin log de accesos/auditoría"],
            ["Continuidad / backup BD", "No cumple", "Definir RPO/RTO"],
        ],
    )
    h2(doc, "14.4 Ley 1581 de 2012 (Colombia)")
    add_table(
        doc,
        ["Dato personal", "Ubicación", "Finalidad"],
        [
            ["Nombre carnicero", "users.nombre, apellido", "Operación"],
            ["Nombre cliente", "pedidos.cliente_nombre", "Identificar pedido"],
            ["Username", "users.username", "Autenticación"],
        ],
    )
    add_table(
        doc,
        ["Obligación", "Estado", "Acción"],
        [
            ["Política de tratamiento", "No cumple", "Elaborar aviso titular (empresa)"],
            ["Medidas de seguridad art. 18", "Parcial", "Hash contraseñas; reforzar HTTPS y CORS"],
            ["DPIA", "No cumple", "Si se amplían datos o integraciones"],
            ["Encargado cloud/hosting", "Parcial", "Contrato con proveedor de hosting"],
        ],
    )

    # 15 Brechas
    h1(doc, "15. Brechas y plan de cierre")
    h2(doc, "Prioridad 1")
    add_table(
        doc,
        ["#", "Brecha", "Marco"],
        [
            ["1", "Política tratamiento datos Ley 1581", "Legal CO"],
            ["2", "Plan de pruebas + automatización (pytest)", "ISO 29119"],
            ["3", "Backup y restauración PostgreSQL", "ISO 27001"],
            ["4", "Restringir CORS y rate limiting en API", "Seguridad"],
            ["5", "Validación de permisos en todos los endpoints", "ISO 27001"],
        ],
    )
    h2(doc, "Prioridad 2")
    add_table(
        doc,
        ["#", "Brecha"],
        [
            ["6", "Auditoría de intentos de login fallidos"],
            ["7", "Recuperación de contraseña"],
            ["8", "CI/CD con análisis de dependencias"],
            ["9", "SLA y monitoreo APM"],
            ["10", "Evaluación accesibilidad WCAG 2.1"],
        ],
    )
    p(
        doc,
        "NOTA: La documentación Word generada anteriormente en el repositorio ProGest (Cañaveral) "
        "corresponde a otro sistema (activos fijos). Este documento (DOC-PEDIDOS-MAYORISTA-001) es el "
        "aplicable exclusivamente al proyecto Pedidos Mayorista.",
    )

    # 16 Glosario
    h1(doc, "16. Glosario y anexos")
    add_table(
        doc,
        ["Término", "Definición"],
        [
            ["Pedido", "Solicitud de preparación de carnes para un cliente en una sede"],
            ["numero_pedido", "Consecutivo visible por sede (#1, #2…)"],
            ["Corte", "Producto cárneo (ej. punta de anca)"],
            ["Tipo de corte", "Presentación: entero, fileteado, molida, etc."],
            ["Sede", "Punto de venta / carnicería"],
            ["Sala Socket.IO", "sede_{id} — canal de eventos por sede"],
            ["Mayorista", "Usuario que captura pedidos del cliente final"],
        ],
    )
    h2(doc, "Anexo A — Códigos HTTP")
    add_table(
        doc,
        ["Código", "Uso en API"],
        [
            ["200", "OK"],
            ["201", "Creado (pedido, sede, etc.)"],
            ["400", "Usuario duplicado, validación"],
            ["401", "Login incorrecto"],
            ["404", "Recurso no encontrado"],
        ],
    )
    h2(doc, "Anexo B — Referencias en repositorio")
    add_table(
        doc,
        ["Archivo", "Contenido"],
        [
            ["README.md", "Guía principal del proyecto"],
            ["README_SETUP.md", "SQL y setup extendido"],
            ["init_db.sql", "Esquema SQL alternativo"],
            [".env.example", "Plantilla variables"],
        ],
    )

    doc.add_page_break()
    h1(doc, "Fin del documento")
    p(doc, f"DOC-PEDIDOS-MAYORISTA-001 — Versión 1.0 — Generado {date.today().isoformat()}")
    p(doc, "© Uso interno Supertiendas Cañaveral.")

    return doc


def main():
    doc = build_document()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Documento generado: {OUT}")
    print(f"Tamaño: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
